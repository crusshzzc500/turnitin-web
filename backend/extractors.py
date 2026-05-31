from __future__ import annotations

import io
import zipfile
from pathlib import Path
from typing import Any

from .integrity import scan_text
from .ocr import extract_pdf_text_with_ocr
from .text import count_words, normalize_display_text


class UnsupportedDocumentError(ValueError):
    pass


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "cp1258", "latin-1"):
        try:
            return normalize_display_text(content.decode(encoding))
        except UnicodeDecodeError:
            continue
    return normalize_display_text(content.decode("utf-8", errors="replace"))


def extract_document(filename: str, content: bytes) -> dict[str, Any]:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".csv", ".rtf"}:
        text = _decode_text(content)
        return {"text": text, "metadata": {"filename": filename}, "integrityFlags": scan_text(text)}
    if suffix == ".pdf":
        return _extract_pdf(filename, content)
    if suffix == ".docx":
        return _extract_docx(filename, content)
    raise UnsupportedDocumentError("Định dạng chưa được hỗ trợ. Hãy dùng .txt, .md, .pdf hoặc .docx.")


def _extract_pdf(filename: str, content: bytes) -> dict[str, Any]:
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise UnsupportedDocumentError("Máy chủ chưa cài thư viện đọc PDF.") from error

    reader = PdfReader(io.BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(pages)
    metadata = {"filename": filename, "pages": len(pages), "pdfMetadata": dict(reader.metadata or {})}
    if count_words(text) < 20:
        ocr_result = extract_pdf_text_with_ocr(content)
        metadata["ocr"] = ocr_result["metadata"]
        if count_words(ocr_result["text"]) > count_words(text):
            text = ocr_result["text"]
    text = normalize_display_text(text)
    return {"text": text, "metadata": metadata, "integrityFlags": scan_text(text)}


def _extract_docx(filename: str, content: bytes) -> dict[str, Any]:
    try:
        from docx import Document
        from docx.shared import RGBColor
    except ImportError as error:
        raise UnsupportedDocumentError("Máy chủ chưa cài thư viện đọc DOCX.") from error

    _validate_docx_archive(content)
    document = Document(io.BytesIO(content))
    paragraphs = [normalize_display_text(paragraph.text) for paragraph in document.paragraphs]
    flags = scan_text("\n".join(paragraphs))
    hidden_runs = 0
    white_runs = 0

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.font.hidden:
                hidden_runs += 1
            if run.font.color and run.font.color.rgb == RGBColor(255, 255, 255):
                white_runs += 1

    if hidden_runs:
        flags.append(
            {
                "kind": "hidden_docx_text",
                "severity": "high",
                "count": hidden_runs,
                "message": "Phát hiện phần văn bản DOCX được đánh dấu ẩn.",
                "details": [],
            }
        )
    if white_runs:
        flags.append(
            {
                "kind": "white_docx_text",
                "severity": "high",
                "count": white_runs,
                "message": "Phát hiện chữ màu trắng có thể khó nhìn thấy.",
                "details": [],
            }
        )

    core = document.core_properties
    metadata = {
        "filename": filename,
        "author": core.author,
        "title": core.title,
        "created": core.created.isoformat() if core.created else None,
        "modified": core.modified.isoformat() if core.modified else None,
        "paragraphs": len(paragraphs),
    }
    return {"text": "\n".join(paragraphs), "metadata": metadata, "integrityFlags": flags}


def _validate_docx_archive(content: bytes, max_expanded_bytes: int = 1_000_000_000) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            expanded_bytes = sum(item.file_size for item in archive.infolist())
    except zipfile.BadZipFile as error:
        raise UnsupportedDocumentError("Tệp DOCX không hợp lệ.") from error
    if expanded_bytes > max_expanded_bytes:
        raise UnsupportedDocumentError("Tệp DOCX giải nén vượt quá giới hạn an toàn.")
