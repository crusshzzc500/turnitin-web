from __future__ import annotations

import json
import base64
import io
import os
import sqlite3
import sys
import tempfile
import threading
import time
import unittest
import zipfile
from contextlib import contextmanager
from dataclasses import replace
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from types import ModuleType, SimpleNamespace
from unittest.mock import Mock, patch
from urllib.error import HTTPError
from urllib.request import Request, urlopen

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))

from backend.analysis import SimilarityAnalyzer
from backend.config import Settings, normalize_gemini_model
from backend.crawler import Crawler
from backend.demo_data import seed_demo_sources
from backend.extractors import extract_document
from backend.jobs import AnalysisJobManager
from backend.search import OpenSearchBackend
from backend.server import create_server
from backend.storage import PostgresConnection, PostgresStorage, Storage, utc_now
from backend.text import normalize_display_text, similarity
from backend.web_discovery import (
    DiscoveryResult,
    WebDiscovery,
    _focused_content_window,
    build_queries,
    build_thorough_queries,
    candidate_relevance,
    normalize_candidate_url,
    regional_coverage,
)
from backend.writing_assistant import CitationWritingAssistant


class LocalSiteHandler(BaseHTTPRequestHandler):
    flaky_hits = 0

    def do_GET(self) -> None:
        base = f"http://{self.headers['Host']}"
        if self.path == "/flaky":
            type(self).flaky_hits += 1
            if type(self).flaky_hits < 3:
                body = "temporary server error"
                payload = body.encode("utf-8")
                self.send_response(503)
                self.send_header("Content-Type", "text/plain; charset=utf-8")
                self.send_header("Content-Length", str(len(payload)))
                self.end_headers()
                self.wfile.write(payload)
                return
        if self.path == "/always-fail":
            body = "persistent server error"
            payload = body.encode("utf-8")
            self.send_response(503)
            self.send_header("Content-Type", "text/plain; charset=utf-8")
            self.send_header("Content-Length", str(len(payload)))
            self.end_headers()
            self.wfile.write(payload)
            return
        pages = {
            "/robots.txt": (
                "text/plain",
                "User-agent: *\nAllow: /allowed\nAllow: /next\nAllow: /flaky\nAllow: /always-fail\nDisallow: /blocked\n",
            ),
            "/allowed": (
                "text/html",
                """
                <html><head><title>Nguồn được phép</title></head>
                <body>
                  <article>
                    <h1>Nguồn công khai</h1>
                    <p>Nội dung công khai đủ dài để crawler lập chỉ mục phục vụ việc đối chiếu tài liệu học thuật.</p>
                    <a href="/next">Trang kế tiếp</a>
                    <a href="/blocked">Trang bị chặn</a>
                  </article>
                </body></html>
                """,
            ),
            "/next": (
                "text/html",
                """
                <html><head><title>Nguồn kế tiếp</title></head>
                <body><p>Trang kế tiếp cũng có nội dung công khai đủ dài để được đưa vào chỉ mục tìm kiếm văn bản.</p></body>
                </html>
                """,
            ),
            "/blocked": (
                "text/html",
                "<html><body><p>Nội dung này không được phép thu thập theo robots.txt và không được lập chỉ mục.</p></body></html>",
            ),
            "/flaky": (
                "text/html",
                """
                <html><head><title>Nguồn hồi phục</title></head>
                <body><p>Nguồn này hồi phục sau lỗi tạm thời và có đủ nội dung để được lập chỉ mục tìm kiếm.</p></body></html>
                """,
            ),
            "/sitemap.xml": (
                "application/xml",
                f"""
                <?xml version="1.0" encoding="UTF-8"?>
                <urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
                  <url><loc>{base}/allowed</loc></url>
                  <url><loc>{base}/next</loc></url>
                </urlset>
                """.strip(),
            ),
        }
        content_type, body = pages.get(self.path, ("text/plain", "not found"))
        status = 200 if self.path in pages else 404
        payload = body.encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", f"{content_type}; charset=utf-8")
        self.send_header("Content-Length", str(len(payload)))
        self.end_headers()
        self.wfile.write(payload)

    def log_message(self, format: str, *args) -> None:
        return


class FakeOpenSearchHandler(BaseHTTPRequestHandler):
    requests: list[dict] = []

    def do_GET(self) -> None:
        self._record_and_send({"cluster_name": "minh-chung-test"})

    def do_PUT(self) -> None:
        self._record_and_send({"acknowledged": True})

    def do_POST(self) -> None:
        if self.path.endswith("/_search"):
            self._record_and_send(
                {
                    "hits": {
                        "hits": [
                            {
                                "_source": {
                                    "chunk_id": 1,
                                    "text_content": "Nội dung OpenSearch thử nghiệm",
                                    "token_count": 4,
                                    "source_id": 9,
                                    "url": "https://example.org/opensearch",
                                    "title": "Nguồn OpenSearch",
                                    "source_type": "website",
                                }
                            }
                        ]
                    }
                }
            )
            return
        self._record_and_send({"acknowledged": True})

    def _record_and_send(self, response: dict) -> None:
        length = int(self.headers.get("Content-Length", "0"))
        body = self.rfile.read(length).decode("utf-8", errors="replace")
        type(self).requests.append({"method": self.command, "path": self.path, "body": body})
        payload = json.dumps(response).encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(payload)))
        self.end_headers()
        self.wfile.write(payload)

    def log_message(self, format: str, *args) -> None:
        return


def settings_for(
    root: Path,
    *,
    port: int = 0,
    allow_private: bool = False,
    max_attempts: int = 3,
    retry_base_seconds: float = 0,
    tavily_api_key: str = "",
    exa_api_key: str = "",
    websearchapi_api_key: str = "",
    linkup_api_key: str = "",
    serper_api_key: str = "",
    brave_search_api_key: str = "",
    gemini_api_key: str = "",
    openai_api_key: str = "",
    public_mode: bool = False,
    auth_mode: str = "demo",
) -> Settings:
    return Settings(
        root_dir=root,
        database_path=root / "test.db",
        static_dir=ROOT,
        host="127.0.0.1",
        port=port,
        crawler_delay_seconds=0,
        crawler_timeout_seconds=3,
        crawler_allow_private_hosts=allow_private,
        crawler_max_attempts=max_attempts,
        crawler_retry_base_seconds=retry_base_seconds,
        tavily_api_key=tavily_api_key,
        exa_api_key=exa_api_key,
        websearchapi_api_key=websearchapi_api_key,
        linkup_api_key=linkup_api_key,
        serper_api_key=serper_api_key,
        brave_search_api_key=brave_search_api_key,
        gemini_api_key=gemini_api_key,
        openai_api_key=openai_api_key,
        public_mode=public_mode,
        auth_mode=auth_mode,
    )


class PostgresStorageTest(unittest.TestCase):
    @staticmethod
    def _modules(connection: Mock) -> dict[str, ModuleType]:
        psycopg = ModuleType("psycopg")
        rows = ModuleType("psycopg.rows")
        psycopg.connect = Mock(return_value=connection)  # type: ignore[attr-defined]
        rows.dict_row = object()  # type: ignore[attr-defined]
        return {"psycopg": psycopg, "psycopg.rows": rows}

    @staticmethod
    def _storage() -> PostgresStorage:
        storage = object.__new__(PostgresStorage)
        storage.database_url = "postgresql://example.invalid/minh_chung"
        storage.search_mirror = None
        storage._initialize_connection_pool()
        return storage

    def test_reuses_healthy_connection(self) -> None:
        connection = Mock(closed=False)
        modules = self._modules(connection)
        with patch.dict(sys.modules, modules):
            storage = self._storage()
            with storage.connect():
                pass
            with storage.connect():
                pass
        modules["psycopg"].connect.assert_called_once()  # type: ignore[attr-defined]
        self.assertEqual(connection.commit.call_count, 2)
        connection.close.assert_not_called()

    def test_discards_connection_when_rollback_fails(self) -> None:
        connection = Mock(closed=False)
        connection.rollback.side_effect = RuntimeError("connection lost")
        modules = self._modules(connection)
        with patch.dict(sys.modules, modules):
            storage = self._storage()
            with self.assertRaisesRegex(ValueError, "bad transaction"):
                with storage.connect():
                    raise ValueError("bad transaction")
        connection.close.assert_called_once()
        self.assertEqual(storage._connection_pool_created, 0)

    def test_postgres_connection_batches_chunk_inserts(self) -> None:
        connection = Mock()
        cursor = connection.cursor.return_value
        wrapper = PostgresConnection(connection)
        wrapper.executemany("INSERT INTO chunks(source_id, position) VALUES (?, ?)", [(1, 0), (1, 1)])
        cursor.executemany.assert_called_once_with(
            "INSERT INTO chunks(source_id, position) VALUES (%s, %s)",
            [(1, 0), (1, 1)],
        )

    def test_postgres_search_uses_full_text_index_query(self) -> None:
        storage = object.__new__(PostgresStorage)
        connection = Mock()
        connection.execute.return_value.fetchall.return_value = []

        @contextmanager
        def connect():
            yield connection

        storage.connect = connect
        storage.search_chunks("Minh bạch dữ liệu giáo dục giúp đối chiếu nguồn tài liệu.", organization_id=1)
        sql, params = connection.execute.call_args.args
        self.assertIn("to_tsvector('simple', chunks.folded_text) @@ to_tsquery('simple', ?)", sql)
        self.assertIn(" | ", params[0])
        self.assertEqual(params[1:], (1, 100))


class SimilarityAnalyzerTest(unittest.TestCase):
    def test_prefetches_document_candidates_before_matching_segments(self) -> None:
        backend = Mock()
        text = (
            "Minh bạch dữ liệu giúp người học kiểm tra nguồn tài liệu một cách rõ ràng. "
            "Quy trình đối chiếu cần ghi nhận đúng nội dung đã xuất hiện trên trang nguồn."
        )
        backend.search_chunks.return_value = [
            {
                "source_id": 1,
                "title": "Nguồn công khai",
                "url": "https://example.org/source",
                "source_type": "website",
                "text_content": text,
            }
        ]
        report = SimilarityAnalyzer(backend).analyze(text)
        self.assertEqual(report["percent"], 100)
        backend.search_chunks.assert_called_once_with(text, limit=500, organization_id=None)

    def test_exact_candidate_skips_expensive_similarity_scoring(self) -> None:
        backend = Mock()
        text = "This complete copied sentence has enough words to be matched exactly against the indexed public source."
        backend.search_chunks.return_value = [
            {
                "source_id": 1,
                "title": "Public source",
                "url": "https://example.org/source",
                "source_type": "website",
                "text_content": f"Intro text. {text} Related text.",
            }
        ]
        with patch("backend.analysis.similarity", side_effect=AssertionError("slow scorer should not run")):
            report = SimilarityAnalyzer(backend).analyze(text)
        self.assertEqual(report["percent"], 100)

    def test_finds_matching_source_and_integrity_flag(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            storage = Storage(Path(directory) / "test.db")
            seed_demo_sources(storage)
            report = SimilarityAnalyzer(storage).analyze(
                "Liêm chính học thuật là nền tảng của một môi trường giáo dục đáng tin cậy. "
                "Người học cần phân biệt rõ việc tham khảo ý tưởng, trích dẫn trực tiếp và sao "
                "chép nội dung mà không ghi nhận nguồn. Ký tự thử nghiệm a\u200bb."
            )
            self.assertGreater(report["percent"], 70)
            self.assertEqual(report["sources"][0]["title"], "Sổ tay về đạo đức học thuật")
            self.assertEqual(report["integrityFlags"][0]["kind"], "zero_width_characters")

    def test_manual_source_is_added_to_search_index(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            storage = Storage(Path(directory) / "test.db")
            storage.upsert_source(
                url="manual://source/test",
                title="Nguồn tự thêm",
                source_type="tự thêm",
                text_content="Nguồn tự thêm có đủ nội dung để được lập chỉ mục và tìm kiếm trong kho dữ liệu.",
            )
            self.assertEqual(storage.stats()["sources"], 1)
            self.assertGreater(storage.stats()["chunks"], 0)

    def test_source_versions_preserve_previous_content(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            storage = Storage(Path(directory) / "test.db")
            source_id = storage.upsert_source(
                url="https://example.org/versioned",
                title="Nguồn phiên bản",
                text_content="Phiên bản đầu tiên có đủ nội dung để được ghi nhận trong lịch sử nguồn dữ liệu.",
            )
            storage.upsert_source(
                url="https://example.org/versioned",
                title="Nguồn phiên bản đã sửa",
                text_content="Phiên bản thứ hai đã thay đổi và vẫn được ghi nhận trong lịch sử nguồn dữ liệu.",
            )
            storage.upsert_source(
                url="https://example.org/versioned",
                title="Nguồn phiên bản đã sửa",
                text_content="Phiên bản thứ hai đã thay đổi và vẫn được ghi nhận trong lịch sử nguồn dữ liệu.",
            )
            versions = storage.list_source_versions(source_id)
            self.assertEqual(storage.stats()["source_versions"], 2)
            self.assertEqual([item["version_number"] for item in versions], [2, 1])
            self.assertEqual(versions[0]["title"], "Nguồn phiên bản đã sửa")

    def test_existing_sources_are_backfilled_with_initial_version(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            database = Path(directory) / "legacy.db"
            connection = sqlite3.connect(database)
            connection.execute(
                """
                CREATE TABLE sources (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    url TEXT NOT NULL UNIQUE,
                    canonical_url TEXT,
                    title TEXT NOT NULL,
                    source_type TEXT NOT NULL DEFAULT 'website',
                    text_content TEXT NOT NULL,
                    content_hash TEXT NOT NULL,
                    word_count INTEGER NOT NULL DEFAULT 0,
                    metadata_json TEXT NOT NULL DEFAULT '{}',
                    fetched_at TEXT NOT NULL,
                    created_at TEXT NOT NULL,
                    updated_at TEXT NOT NULL
                )
                """
            )
            connection.execute(
                """
                INSERT INTO sources(
                    url, title, text_content, content_hash, word_count,
                    fetched_at, created_at, updated_at
                )
                VALUES ('https://example.org/legacy', 'Nguồn cũ', 'Nội dung cũ', 'legacy-hash', 3,
                        '2026-01-01', '2026-01-01', '2026-01-01')
                """
            )
            connection.commit()
            connection.close()
            storage = Storage(database)
            self.assertEqual(storage.stats()["source_versions"], 1)
            self.assertEqual(storage.list_source_versions(1)[0]["title"], "Nguồn cũ")


class ConfigTest(unittest.TestCase):
    def test_render_port_enables_external_bind_and_public_mode(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            with patch.dict(os.environ, {"PORT": "9123"}, clear=True):
                settings = Settings.from_env(Path(directory))
            self.assertEqual(settings.host, "0.0.0.0")
            self.assertEqual(settings.port, 9123)
            self.assertTrue(settings.public_mode)

    def test_platform_public_mode_can_be_explicitly_disabled(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            with patch.dict(os.environ, {"PORT": "9123", "MINH_CHUNG_PUBLIC_MODE": "0"}, clear=True):
                settings = Settings.from_env(Path(directory))
            self.assertFalse(settings.public_mode)

    def test_database_url_enables_password_mode_even_with_stale_public_setting(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            with patch.dict(
                os.environ,
                {
                    "PORT": "9123",
                    "DATABASE_URL": "postgresql://example.invalid/minh_chung",
                    "MINH_CHUNG_PUBLIC_MODE": "1",
                    "MINH_CHUNG_AUTH_MODE": "demo",
                },
                clear=True,
            ):
                settings = Settings.from_env(Path(directory))
            self.assertFalse(settings.public_mode)
            self.assertEqual(settings.auth_mode, "password")

    def test_web_discovery_speed_limits_clamp_stale_environment_values(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            with patch.dict(
                os.environ,
                {
                    "MINH_CHUNG_WEB_DISCOVERY_TIME_BUDGET_SECONDS": "150",
                    "MINH_CHUNG_WEB_DISCOVERY_THOROUGH_TIME_BUDGET_SECONDS": "150",
                    "MINH_CHUNG_WEB_DISCOVERY_REQUEST_TIMEOUT_SECONDS": "45",
                    "MINH_CHUNG_WEB_DISCOVERY_ENRICHMENT_MAX_SOURCES": "4",
                },
                clear=True,
            ):
                settings = Settings.from_env(Path(directory))
            self.assertEqual(settings.web_discovery_time_budget_seconds, 25.0)
            self.assertEqual(settings.web_discovery_thorough_time_budget_seconds, 90.0)
            self.assertEqual(settings.web_discovery_request_timeout_seconds, 8.0)
            self.assertEqual(settings.web_discovery_enrichment_max_sources, 2)

    def test_stale_gemini_model_alias_migrates_to_official_free_model(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            with patch.dict(os.environ, {"GEMINI_MODEL": "gemini-3.5-flash"}, clear=True):
                settings = Settings.from_env(Path(directory))
            self.assertEqual(normalize_gemini_model("gemini-3.5-flash"), "gemini-3-flash-preview")
            self.assertEqual(settings.gemini_model, "gemini-3-flash-preview")


class SearchBackendTest(unittest.TestCase):
    def test_opensearch_adapter_indexes_deletes_searches_and_reports_status(self) -> None:
        FakeOpenSearchHandler.requests = []
        server = ThreadingHTTPServer(("127.0.0.1", 0), FakeOpenSearchHandler)
        thread = threading.Thread(target=server.serve_forever, daemon=True)
        thread.start()
        try:
            backend = OpenSearchBackend(f"http://127.0.0.1:{server.server_address[1]}", "chunks-test")
            backend.replace_source(
                9,
                [
                    {
                        "chunk_id": 1,
                        "text_content": "Nội dung OpenSearch thử nghiệm",
                        "normalized_text": "nội dung opensearch thử nghiệm",
                        "folded_text": "noi dung opensearch thu nghiem",
                        "token_count": 4,
                        "source_id": 9,
                        "url": "https://example.org/opensearch",
                        "title": "Nguồn OpenSearch",
                        "source_type": "website",
                    }
                ],
            )
            matches = backend.search_chunks("Nội dung OpenSearch thử nghiệm")
            status = backend.status()
            paths = [item["path"] for item in FakeOpenSearchHandler.requests]
            self.assertEqual(matches[0]["source_id"], 9)
            self.assertEqual(status["clusterName"], "minh-chung-test")
            self.assertIn("/chunks-test", paths)
            self.assertTrue(any(path.startswith("/chunks-test/_delete_by_query") for path in paths))
            self.assertIn("/_bulk", paths)
            self.assertIn("/chunks-test/_search", paths)
        finally:
            server.shutdown()
            server.server_close()


class ExtractorTest(unittest.TestCase):
    def test_pdf_with_little_embedded_text_uses_optional_ocr_fallback(self) -> None:
        try:
            from pypdf import PdfWriter
        except ImportError:
            self.skipTest("pypdf is not installed")

        stream = io.BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=300)
        writer.write(stream)
        ocr_text = (
            "Nội dung OCR thử nghiệm có đủ số lượng từ để chứng minh rằng PDF scan "
            "được chuyển sang văn bản trước khi tạo báo cáo tương đồng cho người dùng."
        )
        with patch(
            "backend.extractors.extract_pdf_text_with_ocr",
            return_value={"text": ocr_text, "metadata": {"attempted": True, "available": True}},
        ):
            extracted = extract_document("scan.pdf", stream.getvalue())
        self.assertEqual(extracted["text"], ocr_text)
        self.assertTrue(extracted["metadata"]["ocr"]["attempted"])


class WebDiscoveryTest(unittest.TestCase):
    def test_public_web_discovery_serializes_global_budget_state(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            discovery = WebDiscovery(settings_for(root), Storage(root / "test.db"))
            state_lock = threading.Lock()
            active = 0
            maximum_active = 0

            def slow_discovery(*_args, **_kwargs) -> dict:
                nonlocal active, maximum_active
                with state_lock:
                    active += 1
                    maximum_active = max(maximum_active, active)
                time.sleep(0.02)
                with state_lock:
                    active -= 1
                return {"provider": "none"}

            with patch.object(discovery, "_discover_and_index", side_effect=slow_discovery):
                threads = [
                    threading.Thread(
                        target=discovery.discover_and_index,
                        kwargs={"text": "public text", "organization_id": 1},
                    )
                    for _index in range(3)
                ]
                for thread in threads:
                    thread.start()
                for thread in threads:
                    thread.join()
            self.assertEqual(maximum_active, 1)

    def test_gemini_expands_queries_with_structured_output_and_keeps_exact_query(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, gemini_api_key="gemini-test-key"), storage)
            discovery._active_deadline = time.monotonic() + 10
            captured: dict = {}
            initial = [f'"exact phrase {index}"' for index in range(10)]

            def fake_request(url, payload, *, headers, timeout):
                captured.update({"url": url, "payload": payload, "headers": headers, "timeout": timeout})
                return {
                    "candidates": [
                        {
                            "content": {
                                "parts": [
                                    {
                                        "text": json.dumps(
                                            {"queries": ["semantic paraphrase search", "translated original phrase"]}
                                        )
                                    }
                                ]
                            }
                        }
                    ]
                }

            with patch.object(discovery, "_json_request", side_effect=fake_request):
                expanded = discovery._expand_queries_with_gemini("Submitted document text", initial)
            self.assertEqual(
                captured["url"],
                "https://generativelanguage.googleapis.com/v1beta/models/gemini-3-flash-preview:generateContent",
            )
            self.assertEqual(captured["payload"]["generationConfig"]["responseMimeType"], "application/json")
            self.assertEqual(captured["payload"]["generationConfig"]["thinkingConfig"]["thinkingLevel"], "minimal")
            self.assertEqual(
                captured["payload"]["generationConfig"]["responseJsonSchema"]["properties"]["queries"]["maxItems"],
                3,
            )
            self.assertEqual(captured["headers"]["x-goog-api-key"], "gemini-test-key")
            self.assertEqual(expanded[0], initial[0])
            self.assertIn("semantic paraphrase search", expanded)
            self.assertLessEqual(len(expanded), 10)

    def test_gemini_query_expansion_failure_keeps_existing_queries(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, gemini_api_key="gemini-test-key"), storage)
            discovery._active_deadline = time.monotonic() + 10
            initial = ['"exact copied phrase"', "keyword signature"]
            with patch.object(discovery, "_json_request", side_effect=TimeoutError):
                self.assertEqual(discovery._expand_queries_with_gemini("Submitted document text", initial), initial)

    def test_gemini_expansion_is_preferred_before_openai_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(
                settings_for(
                    root,
                    tavily_api_key="tavily",
                    gemini_api_key="gemini",
                    openai_api_key="openai",
                ),
                storage,
            )
            text = "This submitted document contains enough public text to create a useful source discovery query."
            tavily_result = DiscoveryResult("tavily", True, True, ["gemini query"], 0, 0, "No source.", [])
            with (
                patch.object(discovery, "_expand_queries_with_gemini", return_value=["gemini query"]),
                patch.object(discovery, "_expand_queries_with_openai") as openai,
                patch.object(discovery, "_tavily", return_value=tavily_result) as tavily,
            ):
                discovery.discover_and_index(text, organization_id=1)
            openai.assert_not_called()
            self.assertEqual(tavily.call_args.args[0], ["gemini query"])

    def test_openai_expands_queries_with_structured_output_and_keeps_exact_query(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, openai_api_key="openai-test-key"), storage)
            discovery._active_deadline = time.monotonic() + 10
            captured: dict = {}
            initial = [f'"exact phrase {index}"' for index in range(10)]

            def fake_request(url, payload, *, headers, timeout):
                captured.update({"url": url, "payload": payload, "headers": headers, "timeout": timeout})
                return {
                    "output": [
                        {
                            "content": [
                                {
                                    "type": "output_text",
                                    "text": json.dumps(
                                        {"queries": ["semantic paraphrase search", "translated original phrase"]}
                                    ),
                                }
                            ]
                        }
                    ]
                }

            with patch.object(discovery, "_json_request", side_effect=fake_request):
                expanded = discovery._expand_queries_with_openai("Submitted document text", initial)
            self.assertEqual(captured["url"], "https://api.openai.com/v1/responses")
            self.assertEqual(captured["payload"]["model"], "gpt-5-nano")
            self.assertEqual(captured["payload"]["text"]["format"]["type"], "json_schema")
            self.assertEqual(captured["headers"]["Authorization"], "Bearer openai-test-key")
            self.assertEqual(expanded[0], initial[0])
            self.assertIn("semantic paraphrase search", expanded)
            self.assertLessEqual(len(expanded), 10)

    def test_openai_query_expansion_failure_keeps_existing_queries(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, openai_api_key="openai-test-key"), storage)
            discovery._active_deadline = time.monotonic() + 10
            initial = ['"exact copied phrase"', "keyword signature"]
            with patch.object(discovery, "_json_request", side_effect=TimeoutError):
                self.assertEqual(discovery._expand_queries_with_openai("Submitted document text", initial), initial)

    def test_websearchapi_uses_basic_search_without_content_with_server_side_key(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, websearchapi_api_key="websearch-test-key"), storage)
            captured: dict = {}

            def fake_request(url, payload, *, headers, timeout):
                captured.update({"url": url, "payload": payload, "headers": headers, "timeout": timeout})
                return {"organic": []}

            with patch.object(WebDiscovery, "_json_request", side_effect=fake_request):
                discovery._fetch_websearchapi("academic integrity", 10)
            self.assertEqual(captured["url"], "https://api.websearchapi.ai/ai-search")
            self.assertFalse(captured["payload"]["includeContent"])
            self.assertFalse(captured["payload"]["includeAnswer"])
            self.assertEqual(captured["headers"]["Authorization"], "Bearer websearch-test-key")
            self.assertEqual(captured["timeout"], 7.0)

    def test_linkup_uses_fast_search_results_with_server_side_key(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, linkup_api_key="linkup-test-key"), storage)
            captured: dict = {}

            def fake_request(url, payload, *, headers, timeout):
                captured.update({"url": url, "payload": payload, "headers": headers, "timeout": timeout})
                return {"results": []}

            with patch.object(WebDiscovery, "_json_request", side_effect=fake_request):
                discovery._fetch_linkup("academic integrity", 10)
            self.assertEqual(captured["url"], "https://api.linkup.so/v1/search")
            self.assertEqual(captured["payload"]["depth"], "fast")
            self.assertEqual(captured["payload"]["outputType"], "searchResults")
            self.assertEqual(captured["headers"]["Authorization"], "Bearer linkup-test-key")
            self.assertEqual(captured["timeout"], 7.0)

    def test_serper_uses_one_query_shape_with_server_side_key(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, serper_api_key="serper-test-key"), storage)
            captured: dict = {}

            def fake_request(url, payload, *, headers, timeout):
                captured.update({"url": url, "payload": payload, "headers": headers, "timeout": timeout})
                return {"organic": []}

            with patch.object(WebDiscovery, "_json_request", side_effect=fake_request):
                discovery._fetch_serper("academic integrity", 10)
            self.assertEqual(captured["url"], "https://google.serper.dev/search")
            self.assertEqual(captured["payload"], {"q": "academic integrity", "num": 10})
            self.assertEqual(captured["headers"]["X-API-KEY"], "serper-test-key")
            self.assertEqual(captured["timeout"], 7.0)

    def test_serper_prioritizes_relevant_result_before_early_noise(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, serper_api_key="serper"), storage)
            query = '"lòng yêu nước là một trong những giá trị tinh thần cao quý của mỗi con người"'

            def fake_request(*_args, **_kwargs) -> dict:
                return {
                    "organic": [
                        {
                            "link": "https://example.org/noise",
                            "title": "Hình ảnh quê hương",
                            "snippet": "Quê hương có sông suối núi đồi và nhiều phong cảnh thiên nhiên đẹp mắt.",
                        },
                        {
                            "link": "https://example.org/original",
                            "title": "Lòng yêu nước",
                            "snippet": query,
                        },
                    ]
                }

            with patch.object(WebDiscovery, "_json_request", side_effect=fake_request):
                result = discovery._serper([query], organization_id=1, max_results=1)
            self.assertEqual(result.indexed, 1)
            self.assertEqual(result.sources[0]["url"], "https://example.org/original")

    def test_serper_hard_caps_queries_even_if_called_with_more_than_one(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, serper_api_key="serper-test-key"), storage)
            with patch.object(discovery, "_fetch_serper", return_value={"organic": []}) as fetch:
                result = discovery._serper(["one", "two", "three"], organization_id=1, max_results=10)
            self.assertEqual(result.queries, ["one"])
            fetch.assert_called_once_with("one", 10)

    def test_new_fallbacks_hard_cap_queries_even_if_called_with_more_than_one(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(
                settings_for(root, websearchapi_api_key="websearch", linkup_api_key="linkup"),
                storage,
            )
            with (
                patch.object(discovery, "_fetch_websearchapi", return_value={"organic": []}) as websearch,
                patch.object(discovery, "_fetch_linkup", return_value={"results": []}) as linkup,
            ):
                websearch_result = discovery._websearchapi(["one", "two"], organization_id=1, max_results=10)
                linkup_result = discovery._linkup(["one", "two"], organization_id=1, max_results=10)
            self.assertEqual(websearch_result.queries, ["one"])
            self.assertEqual(linkup_result.queries, ["one"])
            websearch.assert_called_once_with("one", 10)
            linkup.assert_called_once_with("one", 10)

    def test_exa_uses_instant_search_with_highlights_and_server_side_key(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, exa_api_key="exa-test-key"), storage)
            captured: dict = {}

            def fake_request(_url, payload, *, headers, timeout):
                captured.update({"payload": payload, "headers": headers, "timeout": timeout})
                return {"results": []}

            with patch.object(WebDiscovery, "_json_request", side_effect=fake_request):
                discovery._fetch_exa("liêm chính học thuật", 10)
            self.assertEqual(captured["payload"]["type"], "instant")
            self.assertEqual(captured["payload"]["numResults"], 10)
            self.assertEqual(captured["payload"]["contents"]["highlights"]["maxCharacters"], 1200)
            self.assertEqual(captured["headers"]["x-api-key"], "exa-test-key")
            self.assertEqual(captured["timeout"], 7.0)

    def test_exa_fallback_runs_only_when_tavily_returns_too_few_sources(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, tavily_api_key="tavily", exa_api_key="exa"), storage)
            text = (
                "Đoạn đầu tiên có đủ số lượng từ và chứa nội dung riêng để dùng làm truy vấn tìm kiếm công khai. "
                "Đoạn thứ hai cũng có đủ số lượng từ nhưng mang thêm nhiều thuật ngữ học thuật khác nhau để kiểm tra. "
                "Đoạn thứ ba tiếp tục bổ sung nội dung nhằm chắc chắn bộ chọn không gửi toàn bộ tài liệu ra bên ngoài."
            )
            primary = DiscoveryResult("tavily", True, True, ["primary"], 0, 0, "Tavily thiếu nguồn.", [])
            fallback = DiscoveryResult(
                "exa",
                True,
                True,
                ["fallback"],
                1,
                0,
                "Exa bổ sung nguồn.",
                [{"id": 9, "title": "Nguồn Exa", "url": "https://example.org/exa"}],
            )
            with (
                patch.object(discovery, "_tavily", return_value=primary),
                patch.object(discovery, "_exa", return_value=fallback) as exa,
            ):
                result = discovery.discover_and_index(text, organization_id=1)
            self.assertEqual(result["provider"], "tavily+exa")
            self.assertEqual(result["indexed"], 1)
            self.assertLessEqual(len(exa.call_args.args[0]), 3)

    def test_exa_fallback_is_skipped_when_tavily_has_enough_sources(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(
                settings_for(root, tavily_api_key="tavily", exa_api_key="exa"),
                storage,
            )
            text = "Nội dung đủ dài để tạo truy vấn kiểm tra và xác minh bộ điều phối không gọi Exa khi Tavily đã đủ nguồn."
            sources = [
                {"id": index, "title": f"Nguồn {index}", "url": f"https://example.org/{index}"}
                for index in range(8)
            ]
            primary = DiscoveryResult("tavily", True, True, ["primary"], 8, 0, "Tavily đủ nguồn.", sources)
            with (
                patch.object(discovery, "_tavily", return_value=primary),
                patch.object(discovery, "_exa") as exa,
            ):
                result = discovery.discover_and_index(text, organization_id=1)
            self.assertEqual(result["provider"], "tavily")
            exa.assert_not_called()

    def test_serper_precision_runs_before_tavily_and_exa(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(
                settings_for(root, tavily_api_key="tavily", exa_api_key="exa", serper_api_key="serper"),
                storage,
            )
            text = "This sufficiently long excerpt exists to verify the final public web search fallback provider."
            tavily = DiscoveryResult("tavily", True, True, ["primary"], 0, 0, "Tavily empty.", [])
            exa = DiscoveryResult(
                "exa",
                True,
                True,
                ["fallback-exa"],
                1,
                0,
                "Exa added one source.",
                [{"id": 9, "title": "Exa source", "url": "https://example.org/exa"}],
            )
            serper_result = DiscoveryResult(
                "serper",
                True,
                True,
                ["fallback-serper"],
                1,
                0,
                "Serper added one source.",
                [{"id": 10, "title": "Serper source", "url": "https://example.org/serper"}],
            )
            calls: list[str] = []
            with (
                patch.object(discovery, "_tavily", side_effect=lambda *args, **kwargs: (calls.append("tavily"), tavily)[1]) as tavily_search,
                patch.object(discovery, "_exa", side_effect=lambda *args, **kwargs: (calls.append("exa"), exa)[1]),
                patch.object(discovery, "_serper", side_effect=lambda *args, **kwargs: (calls.append("serper"), serper_result)[1]) as serper,
            ):
                result = discovery.discover_and_index(text, organization_id=1)
            self.assertEqual(calls, ["serper", "tavily", "exa"])
            self.assertEqual(result["provider"], "serper+tavily+exa")
            self.assertEqual(result["indexed"], 2)
            self.assertLessEqual(len(serper.call_args.args[0]), 1)
            self.assertEqual(tavily_search.call_args.kwargs["initial_seen_urls"], {"https://example.org/serper"})

    def test_exact_serper_match_skips_broad_fallback_search(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(
                settings_for(root, tavily_api_key="tavily", serper_api_key="serper", gemini_api_key="gemini"),
                storage,
            )
            text = "This complete copied document is long enough to verify that broad fallback search stops after an exact public match."
            serper_result = DiscoveryResult(
                "serper",
                True,
                True,
                ["precision"],
                1,
                0,
                "Serper found the exact source.",
                [{"id": 1, "title": "Exact source", "url": "https://example.org/exact", "exactDocumentMatch": True}],
            )
            with (
                patch.object(discovery, "_serper", return_value=serper_result),
                patch.object(discovery, "_expand_queries_with_gemini") as gemini,
                patch.object(discovery, "_tavily") as tavily,
            ):
                result = discovery.discover_and_index(text, organization_id=1)
            self.assertEqual(result["provider"], "serper")
            gemini.assert_not_called()
            tavily.assert_not_called()

    def test_new_fallbacks_run_between_exa_and_serper_when_sources_are_still_missing(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(
                settings_for(
                    root,
                    tavily_api_key="tavily",
                    exa_api_key="exa",
                    websearchapi_api_key="websearch",
                    linkup_api_key="linkup",
                    serper_api_key="serper",
                ),
                storage,
            )
            text = "This sufficiently long excerpt verifies the complete quota saving public web fallback chain."
            calls: list[str] = []

            def result(provider: str, url: str) -> DiscoveryResult:
                return DiscoveryResult(
                    provider,
                    True,
                    True,
                    [provider],
                    1,
                    0,
                    f"{provider} source.",
                    [{"id": len(calls), "title": provider, "url": url}],
                )

            with (
                patch.object(discovery, "_tavily", side_effect=lambda *args, **kwargs: (calls.append("tavily"), result("tavily", "https://example.org/tavily"))[1]),
                patch.object(discovery, "_exa", side_effect=lambda *args, **kwargs: (calls.append("exa"), result("exa", "https://example.org/exa"))[1]),
                patch.object(discovery, "_websearchapi", side_effect=lambda *args, **kwargs: (calls.append("websearchapi"), result("websearchapi", "https://example.org/websearchapi"))[1]) as websearch,
                patch.object(discovery, "_linkup", side_effect=lambda *args, **kwargs: (calls.append("linkup"), result("linkup", "https://example.org/linkup"))[1]) as linkup,
                patch.object(discovery, "_serper", side_effect=lambda *args, **kwargs: (calls.append("serper"), result("serper", "https://example.org/serper"))[1]) as serper,
            ):
                result_payload = discovery.discover_and_index(text, organization_id=1)
            self.assertEqual(calls, ["serper", "tavily", "exa", "websearchapi", "linkup"])
            self.assertEqual(result_payload["provider"], "serper+tavily+exa+websearchapi+linkup")
            self.assertLessEqual(len(websearch.call_args.args[0]), 1)
            self.assertLessEqual(len(linkup.call_args.args[0]), 1)
            self.assertLessEqual(len(serper.call_args.args[0]), 1)

    def test_thorough_verification_consults_each_configured_provider_with_small_caps(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(
                settings_for(
                    root,
                    tavily_api_key="tavily",
                    exa_api_key="exa",
                    websearchapi_api_key="websearch",
                    linkup_api_key="linkup",
                    serper_api_key="serper",
                    brave_search_api_key="brave",
                ),
                storage,
            )
            text = "This sufficiently long excerpt verifies thorough public web source discovery with bounded provider quotas."
            calls: list[str] = []

            def result(provider: str, count: int) -> DiscoveryResult:
                return DiscoveryResult(
                    provider,
                    True,
                    True,
                    [provider],
                    count,
                    0,
                    f"{provider} sources.",
                    [
                        {
                            "id": f"{provider}-{index}",
                            "title": provider,
                            "url": f"https://example.org/{provider}/{index}",
                        }
                        for index in range(count)
                    ],
                )

            with (
                patch.object(discovery, "_serper", side_effect=lambda *args, **kwargs: (calls.append("serper"), result("serper", 1))[1]),
                patch.object(discovery, "_tavily", side_effect=lambda *args, **kwargs: (calls.append("tavily"), result("tavily", 6))[1]) as tavily,
                patch.object(discovery, "_exa", side_effect=lambda *args, **kwargs: (calls.append("exa"), result("exa", 4))[1]) as exa,
                patch.object(discovery, "_websearchapi", side_effect=lambda *args, **kwargs: (calls.append("websearchapi"), result("websearchapi", 3))[1]) as websearch,
                patch.object(discovery, "_linkup", side_effect=lambda *args, **kwargs: (calls.append("linkup"), result("linkup", 3))[1]) as linkup,
                patch.object(discovery, "_brave", side_effect=lambda *args, **kwargs: (calls.append("brave"), result("brave", 3))[1]) as brave,
            ):
                payload = discovery.discover_and_index(text, organization_id=1, max_results=20, thorough=True)
            self.assertEqual(calls, ["serper", "tavily", "exa", "websearchapi", "linkup", "brave"])
            self.assertEqual(payload["verificationMode"], "thorough")
            self.assertEqual(payload["queryStrategy"], "whole-document-fingerprint-v4")
            self.assertEqual(payload["indexed"], 20)
            self.assertEqual(tavily.call_args.args[2], 6)
            self.assertEqual(exa.call_args.args[2], 4)
            self.assertEqual(websearch.call_args.args[2], 3)
            self.assertEqual(linkup.call_args.args[2], 3)
            self.assertEqual(brave.call_args.args[2], 3)

    def test_tavily_uses_fast_snippets_without_raw_content(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, tavily_api_key="test-key"), storage)
            captured: dict = {}

            def fake_request(_url, payload, *, headers, timeout):
                captured.update({"payload": payload, "headers": headers, "timeout": timeout})
                return {"results": []}

            with patch.object(WebDiscovery, "_json_request", side_effect=fake_request):
                discovery._fetch_tavily("liêm chính học thuật", 10)
            self.assertEqual(captured["payload"]["search_depth"], "fast")
            self.assertFalse(captured["payload"]["include_raw_content"])
            self.assertEqual(captured["timeout"], 7.0)

    def test_build_queries_selects_a_small_number_of_excerpts(self) -> None:
        text = (
            "Đoạn đầu tiên có đủ số lượng từ và chứa nội dung riêng để dùng làm truy vấn tìm kiếm công khai. "
            "Đoạn thứ hai cũng có đủ số lượng từ nhưng mang thêm nhiều thuật ngữ học thuật khác nhau để kiểm tra. "
            "Đoạn thứ ba tiếp tục bổ sung nội dung nhằm chắc chắn bộ chọn không gửi toàn bộ tài liệu ra bên ngoài."
        )
        queries = build_queries(text, max_queries=2)
        self.assertEqual(len(queries), 2)
        self.assertTrue(all(len(query) <= 360 for query in queries))

    def test_build_queries_prioritizes_quoted_exact_phrases(self) -> None:
        text = (
            "Lòng yêu nước là một trong những giá trị tinh thần cao quý của mỗi con người và là trách nhiệm đối với quê hương. "
            "Mỗi công dân cần thể hiện tinh thần đó bằng hành động cụ thể trong học tập, lao động và bảo vệ môi trường. "
            "Sự đoàn kết giúp cộng đồng vượt qua thử thách để cùng xây dựng đất nước ngày càng phát triển."
        )
        queries = build_queries(text, max_queries=3)
        self.assertEqual(len(queries), 3)
        self.assertTrue(all(query.startswith('"') and query.endswith('"') for query in queries))

    def test_build_queries_adds_diverse_keyword_fingerprints_without_extra_quota(self) -> None:
        text = (
            "Khung quản trị dữ liệu giáo dục cần minh bạch trách nhiệm giải trình và bảo vệ quyền riêng tư người học. "
            "Khung quản trị dữ liệu giáo dục cần minh bạch trách nhiệm giải trình và bảo vệ quyền riêng tư người học trong trường học. "
            "Mô hình kiểm định độc lập bổ sung truy vết nguồn tài liệu cùng bằng chứng học thuật có thể xác minh. "
            "Quy trình rà soát cuối cùng phân loại trích dẫn hợp lệ và nội dung cần bổ sung nguồn tham khảo."
        )
        queries = build_queries(text, max_queries=4)
        self.assertEqual(len(queries), 4)
        self.assertTrue(any(len(query.split()) < 10 for query in queries))
        self.assertLess(
            sum("Khung quản trị dữ liệu giáo dục" in query for query in queries),
            2,
        )

    def test_build_queries_splits_long_sentences_into_searchable_windows(self) -> None:
        text = " ".join(f"thuậtngữ{index}" for index in range(52)) + "."
        queries = build_queries(text, max_queries=3)
        self.assertGreaterEqual(len(queries), 2)
        self.assertTrue(all(10 <= len(query.split()) <= 32 for query in queries))

    def test_thorough_queries_cover_beginning_middle_and_end_of_long_document(self) -> None:
        text = " ".join(
            [
                "đầubài minh chứng nguồn công khai cần được kiểm tra bằng truy vấn chính xác xuyên suốt tài liệu.",
                "vùnghai dữ liệu giáo dục minh bạch giúp người học rà soát nội dung tham khảo có trách nhiệm.",
                "vùngba bằng chứng đối chiếu đáng tin cậy cần liên kết đúng trang nguồn đã được xác minh.",
                "giữabài quy trình học thuật phân biệt trích dẫn hợp lệ với nội dung sao chép chưa ghi nhận nguồn.",
                "vùngnăm hệ thống tìm kiếm lập chỉ mục đoạn văn phù hợp để báo cáo có thể giải trình rõ ràng.",
                "vùngsáu người đọc cần kiểm tra từng ghi chú nguồn trước khi sử dụng bản đề xuất chỉnh sửa.",
                "cuốibài dấu vân tay nội dung giúp phát hiện đoạn sao chép nằm xa phần mở đầu của tài liệu.",
            ]
        )
        queries = build_thorough_queries(text, max_queries=7)
        joined = " ".join(queries)
        self.assertLessEqual(len(queries), 7)
        self.assertIn("đầubài", joined)
        self.assertIn("giữabài", joined)
        self.assertIn("cuốibài", joined)

    def test_regional_coverage_distinguishes_searches_from_verified_url_evidence(self) -> None:
        regions = [
            "beginmarker alpha01 alpha02 alpha03 alpha04 alpha05 alpha06 alpha07 alpha08 alpha09 alpha10 alpha11 alpha12 alpha13 alpha14 alpha15 alpha16 alpha17 alpha18 alpha19.",
            "secondmarker beta01 beta02 beta03 beta04 beta05 beta06 beta07 beta08 beta09 beta10 beta11 beta12 beta13 beta14 beta15 beta16 beta17 beta18 beta19.",
            "middlemarker gamma01 gamma02 gamma03 gamma04 gamma05 gamma06 gamma07 gamma08 gamma09 gamma10 gamma11 gamma12 gamma13 gamma14 gamma15 gamma16 gamma17 gamma18 gamma19.",
            "fourthmarker delta01 delta02 delta03 delta04 delta05 delta06 delta07 delta08 delta09 delta10 delta11 delta12 delta13 delta14 delta15 delta16 delta17 delta18 delta19.",
            "endmarker omega01 omega02 omega03 omega04 omega05 omega06 omega07 omega08 omega09 omega10 omega11 omega12 omega13 omega14 omega15 omega16 omega17 omega18 omega19.",
        ]
        coverage = regional_coverage(
            " ".join(regions),
            fingerprint_queries=regions,
            searched_queries=regions,
            sources=[
                {"matchedQuery": regions[0]},
                {"matchedQuery": regions[-1]},
            ],
        )
        self.assertEqual(coverage["totalRegions"], 5)
        self.assertEqual(coverage["fingerprintedRegions"], 5)
        self.assertEqual(coverage["searchedRegions"], 5)
        self.assertGreaterEqual(coverage["evidenceRegions"], 2)
        self.assertLess(coverage["evidenceRegions"], 5)
        self.assertTrue(coverage["needsReview"])
        self.assertIn("khong chung minh", coverage["warning"])

    def test_exact_document_match_marks_all_regions_as_verified(self) -> None:
        coverage = regional_coverage(
            " ".join(f"documentword{index}" for index in range(100)),
            fingerprint_queries=[],
            searched_queries=[],
            sources=[{"exactDocumentMatch": True}],
        )
        self.assertEqual(coverage["totalRegions"], 5)
        self.assertEqual(coverage["searchedRegions"], 5)
        self.assertEqual(coverage["evidenceRegions"], 5)
        self.assertTrue(coverage["completeDocumentMatch"])
        self.assertFalse(coverage["needsReview"])

    def test_thorough_discovery_sends_whole_document_queries_to_provider(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            discovery = WebDiscovery(settings_for(root, tavily_api_key="tavily"), Storage(root / "test.db"))
            text = " ".join(
                [
                    "BEGINMARKER verified public evidence should be searched with precise queries across the document.",
                    "regiontwo academic review records sources carefully before a similarity report is accepted.",
                    "regionthree transparent attribution helps readers inspect the supporting public source page.",
                    "MIDDLEMARKER responsible citation workflow distinguishes quotations from uncited copied passages.",
                    "regionfive similarity evidence remains linked to the indexed canonical public source address.",
                    "regionsix reviewers inspect every citation marker before using an assisted revision proposal.",
                    "ENDMARKER document fingerprints detect copied passages located far from the opening paragraph.",
                ]
            )
            empty = DiscoveryResult("tavily", True, True, [], 0, 0, "No matching source.", [])
            with patch.object(discovery, "_tavily", return_value=empty) as tavily:
                payload = discovery.discover_and_index(text, organization_id=1, thorough=True)
            submitted_queries = " ".join(tavily.call_args.args[0])
            self.assertIn("BEGINMARKER", submitted_queries)
            self.assertIn("MIDDLEMARKER", submitted_queries)
            self.assertIn("ENDMARKER", submitted_queries)
            self.assertEqual(payload["queryStrategy"], "whole-document-fingerprint-v4")

    def test_candidate_relevance_rewards_phrase_overlap_and_rejects_noise(self) -> None:
        query = "quản trị dữ liệu giáo dục minh bạch trách nhiệm giải trình"
        relevant = candidate_relevance(
            query,
            "Quản trị dữ liệu giáo dục",
            "Khung quản trị dữ liệu giáo dục minh bạch trách nhiệm giải trình giúp bảo vệ người học.",
        )
        unrelated = candidate_relevance(
            query,
            "Bản tin thể thao",
            "Lịch thi đấu mới nhất được cập nhật sau vòng chung kết cuối tuần.",
        )
        self.assertGreater(relevant, 0.8)
        self.assertEqual(unrelated, 0.0)

    def test_focused_content_window_keeps_anchor_without_indexing_entire_page(self) -> None:
        anchor = "lòng yêu nước là một phẩm chất quý báu"
        page = f"{'phần giới thiệu ' * 120}{anchor}. {'phần kết luận ' * 120}"
        focused = _focused_content_window(page, [anchor], maximum_chars=500)
        self.assertLessEqual(len(focused), 500)
        self.assertIn(anchor, focused)

    def test_focused_content_window_prefers_complete_submitted_article(self) -> None:
        article = " ".join(f"article-word-{index}" for index in range(80))
        page = f"{'navigation ' * 200}{article}{'related-article ' * 200}"
        focused = _focused_content_window(page, [article], maximum_chars=4_000)
        self.assertEqual(focused, article)

    def test_tracking_urls_are_normalized_before_indexing(self) -> None:
        self.assertEqual(
            normalize_candidate_url("HTTPS://Example.org/article/?utm_source=test&fbclid=123#section"),
            "https://example.org/article",
        )
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root), storage)
            seen_urls: set[str] = set()
            content = (
                "Khung quản trị dữ liệu giáo dục minh bạch trách nhiệm giải trình giúp bảo vệ quyền riêng tư "
                "và nâng cao khả năng truy vết nguồn tài liệu trong môi trường học thuật."
            )
            first = discovery._index_candidate(
                provider="test",
                query="quản trị dữ liệu giáo dục minh bạch trách nhiệm giải trình",
                canonical_url="https://example.org/article?utm_source=first",
                title="Khung quản trị dữ liệu giáo dục",
                content=content,
                organization_id=1,
                minimum_words=8,
                seen_urls=seen_urls,
            )
            duplicate = discovery._index_candidate(
                provider="test",
                query="quản trị dữ liệu giáo dục minh bạch trách nhiệm giải trình",
                canonical_url="https://EXAMPLE.org/article/?fbclid=second",
                title="Bản sao tracking",
                content=content,
                organization_id=1,
                minimum_words=8,
                seen_urls=seen_urls,
            )
            noise = discovery._index_candidate(
                provider="test",
                query="quản trị dữ liệu giáo dục minh bạch trách nhiệm giải trình",
                canonical_url="https://example.org/unrelated",
                title="Bản tin thể thao",
                content="Lịch thi đấu mới nhất được cập nhật sau vòng chung kết cuối tuần với nhiều thay đổi đáng chú ý.",
                organization_id=1,
                minimum_words=8,
                seen_urls=seen_urls,
            )
            self.assertTrue(first)
            self.assertIsNone(duplicate)
            self.assertIsNone(noise)
            self.assertEqual(storage.stats(1)["sources"], 1)

    def test_short_search_snippet_is_enriched_from_allowed_public_page(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root), storage)

            class FakeCrawler:
                url_policy = SimpleNamespace(validate=lambda url: url)
                robots = SimpleNamespace(allowed=lambda _url: True)

                @staticmethod
                def _fetch(_url):
                    return SimpleNamespace(
                        text=(
                            "Khung quản trị dữ liệu giáo dục minh bạch trách nhiệm giải trình giúp bảo vệ quyền riêng tư "
                            "người học trong môi trường học thuật. Nội dung đầy đủ bổ sung bằng chứng truy vết nguồn tài liệu "
                            "và giải thích rõ quy trình kiểm định độc lập cho từng báo cáo nghiên cứu."
                        )
                    )

            discovery.attach_crawler(FakeCrawler())
            discovery._enrichment_remaining = 1
            indexed = discovery._index_candidate(
                provider="test",
                query="quản trị dữ liệu giáo dục minh bạch trách nhiệm giải trình",
                canonical_url="https://example.org/full-article",
                title="Quản trị dữ liệu giáo dục",
                content="Quản trị dữ liệu giáo dục minh bạch trách nhiệm giải trình bảo vệ người học.",
                organization_id=1,
                minimum_words=8,
                seen_urls=set(),
            )
            with storage.connect() as connection:
                row = connection.execute(
                    "SELECT text_content, metadata_json FROM sources WHERE id = ?",
                    (indexed["id"],),
                ).fetchone()
            self.assertIn("Nội dung đầy đủ", row["text_content"])
            self.assertTrue(json.loads(row["metadata_json"])["enrichedFromPublicPage"])

    def test_precision_source_fetches_full_page_even_when_search_excerpt_is_long(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root), storage)
            fetched: list[str] = []
            submitted_text = "submitted exact phrase for the original public article"

            class FakeCrawler:
                url_policy = SimpleNamespace(validate=lambda url: url)
                robots = SimpleNamespace(allowed=lambda _url: True)

                @staticmethod
                def _fetch(url):
                    fetched.append(url)
                    return SimpleNamespace(
                        text=f"{'intro ' * 150}{submitted_text} {'complete-source-marker ' * 180}"
                    )

            discovery.attach_crawler(FakeCrawler())
            discovery._enrichment_remaining = 1
            content, enriched = discovery._enrich_content(
                "https://example.org/original",
                "search-excerpt " * 150,
                relevance=0.95,
                query=submitted_text,
                comparison_text=submitted_text,
            )
            self.assertTrue(enriched)
            self.assertEqual(fetched, ["https://example.org/original"])
            self.assertIn("complete-source-marker", content)

    def test_tavily_results_are_namespaced_per_organization(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, tavily_api_key="test-key"), storage)
            content = (
                "Nguồn Tavily công khai có nội dung đủ dài để được lập chỉ mục cho việc kiểm tra tương đồng. "
                "Dữ liệu này lặp lại một số cụm từ mô phỏng bài viết học thuật và giúp xác minh cách lưu nguồn "
                "riêng theo từng tổ chức mà không ghi đè hoặc làm rò rỉ dữ liệu giữa hai trường khác nhau. "
                "Phần bổ sung bảo đảm văn bản vượt qua ngưỡng tối thiểu bốn mươi từ của lớp tìm kiếm web."
            )
            response = {"results": [{"url": "https://example.org/shared", "title": "Nguồn chung", "raw_content": content}]}
            with storage.connect() as connection:
                cursor = connection.execute(
                    "INSERT INTO organizations(slug, name, created_at) VALUES (?, ?, ?)",
                    ("other-school", "Trường Khác", utc_now()),
                )
                second_organization_id = int(cursor.lastrowid)
            with patch.object(WebDiscovery, "_json_request", return_value=response):
                first = discovery.discover_and_index(content, organization_id=1)
                second = discovery.discover_and_index(content, organization_id=second_organization_id)
            self.assertEqual(first["indexed"], 1)
            self.assertEqual(second["indexed"], 1)
            self.assertNotEqual(first["sources"][0]["id"], second["sources"][0]["id"])
            with storage.connect() as connection:
                urls = {
                    row["organization_id"]: row["url"]
                    for row in connection.execute(
                        "SELECT organization_id, url FROM sources WHERE source_type = 'web-tavily'"
                    )
                }
            self.assertTrue(urls[1].startswith("web-discovery://1/"))
            self.assertTrue(urls[second_organization_id].startswith(f"web-discovery://{second_organization_id}/"))

    def test_tavily_queries_run_in_parallel_and_report_progress(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            settings = replace(
                settings_for(root, tavily_api_key="test-key"),
                web_discovery_max_queries=3,
                web_discovery_parallel_workers=3,
            )
            discovery = WebDiscovery(settings, storage)
            text = (
                "Đoạn đầu tiên có đủ số lượng từ riêng biệt để làm truy vấn tìm kiếm công khai cho tài liệu. "
                "Đoạn thứ hai bổ sung nhiều thuật ngữ học thuật khác nhau nhằm kiểm tra xử lý tìm nguồn đồng thời. "
                "Đoạn thứ ba tiếp tục cung cấp nội dung độc lập để xác minh tiến trình quét web được cập nhật."
            )
            active = 0
            max_active = 0
            lock = threading.Lock()
            updates: list[tuple[int, int, int]] = []

            def fake_request(*_args, **_kwargs) -> dict:
                nonlocal active, max_active
                with lock:
                    active += 1
                    max_active = max(max_active, active)
                try:
                    time.sleep(0.08)
                    return {"results": []}
                finally:
                    with lock:
                        active -= 1

            with patch.object(WebDiscovery, "_json_request", side_effect=fake_request):
                result = discovery.discover_and_index(
                    text,
                    organization_id=1,
                    progress_callback=lambda completed, total, indexed: updates.append((completed, total, indexed)),
                )
            self.assertGreaterEqual(max_active, 2)
            self.assertEqual(result["provider"], "tavily")
            self.assertEqual(updates[-1][:2], (3, 3))

    def test_tavily_prefers_best_excerpt_before_deduplicating_url(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, tavily_api_key="test-key"), storage)
            query = "academic integrity requires transparent citation records for every submitted research document"
            weak = "academic integrity citation records with unrelated filler words for a public page search result"
            strong = f"{query} and preserves the complete original public source content for comparison"
            response = {
                "results": [
                    {"url": "https://example.org/shared", "title": "Weak excerpt", "content": weak},
                    {"url": "https://example.org/shared", "title": "Strong excerpt", "content": strong},
                ]
            }
            with patch.object(discovery, "_fetch_tavily", return_value=response):
                result = discovery._tavily([query], organization_id=1, max_results=10)
            with storage.connect() as connection:
                row = connection.execute("SELECT text_content FROM sources").fetchone()
            self.assertEqual(result.indexed, 1)
            self.assertEqual(row["text_content"], strong)

    def test_tavily_caps_indexed_sources_across_all_queries(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            settings = replace(
                settings_for(root, tavily_api_key="test-key"),
                web_discovery_max_queries=3,
                web_discovery_parallel_workers=3,
            )
            discovery = WebDiscovery(settings, storage)
            text = (
                "Đoạn thứ nhất có đủ số lượng từ riêng biệt để tạo truy vấn tìm kiếm công khai và kiểm tra giới hạn nguồn toàn báo cáo. "
                "Đoạn thứ hai bổ sung thuật ngữ học thuật khác nhau để xác minh nhiều truy vấn không ghi quá nhiều nguồn vào dữ liệu. "
                "Đoạn thứ ba tiếp tục cung cấp nội dung độc lập nhằm bảo đảm báo cáo được trả về nhanh sau khi tìm đủ nguồn phù hợp."
            )
            sequence = 0
            lock = threading.Lock()

            def fake_request(_url, payload, **_kwargs) -> dict:
                nonlocal sequence
                with lock:
                    start = sequence
                    sequence += 20
                query = payload["query"]
                return {
                    "results": [
                        {
                            "url": f"https://example.org/source-{start + index}",
                            "title": query,
                            "content": f"{query} Nội dung bổ sung đủ dài để lập chỉ mục nguồn công khai phù hợp cho báo cáo kiểm tra.",
                        }
                        for index in range(20)
                    ]
                }

            with patch.object(WebDiscovery, "_json_request", side_effect=fake_request):
                result = discovery.discover_and_index(text, organization_id=1, max_results=10)
            self.assertEqual(result["indexed"], 10)
            self.assertEqual(storage.stats(1)["sources"], 10)

    def test_low_relevance_source_does_not_consume_enrichment_budget(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            settings = replace(
                settings_for(root),
                web_discovery_enrichment_max_sources=1,
            )
            discovery = WebDiscovery(settings, storage)
            fetched: list[str] = []

            class FakeCrawler:
                url_policy = SimpleNamespace(validate=lambda url: url)
                robots = SimpleNamespace(allowed=lambda _url: True)

                @staticmethod
                def _fetch(url):
                    fetched.append(url)
                    return SimpleNamespace(text="Nội dung đầy đủ từ trang nguồn chính có thêm nhiều từ để lập chỉ mục.")

            discovery.attach_crawler(FakeCrawler())
            discovery._enrichment_remaining = 1
            discovery._enrich_content("https://example.org/noise", "Nội dung ngắn.", relevance=0.40)
            discovery._enrich_content("https://example.org/original", "Nội dung ngắn.", relevance=0.90)
            self.assertEqual(fetched, ["https://example.org/original"])

    def test_discovered_web_source_is_used_by_similarity_report(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            discovery = WebDiscovery(settings_for(root, tavily_api_key="test-key"), storage)
            text = (
                "Minh bạch dữ liệu giáo dục giúp người học kiểm tra nguồn tài liệu và hiểu rõ trách nhiệm giải trình. "
                "Quy trình đối chiếu công khai cần ghi nhận đúng đoạn văn đã xuất hiện trên trang nguồn."
            )

            def fake_request(_url, payload, **_kwargs) -> dict:
                return {
                    "results": [
                        {
                            "url": "https://example.org/copied-article",
                            "title": "Minh bạch dữ liệu giáo dục",
                            "content": text,
                        }
                    ]
                }

            with patch.object(WebDiscovery, "_json_request", side_effect=fake_request):
                discovered = discovery.discover_and_index(text, organization_id=1)
            report = SimilarityAnalyzer(storage).analyze(text, organization_id=1)
            self.assertEqual(discovered["indexed"], 1)
            self.assertEqual(report["percent"], 100)
            self.assertEqual(report["sources"][0]["url"], "https://example.org/copied-article")

    def test_tavily_returns_partial_results_after_time_budget(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            settings = replace(
                settings_for(root, tavily_api_key="test-key"),
                web_discovery_max_queries=3,
                web_discovery_parallel_workers=3,
                web_discovery_time_budget_seconds=0.03,
            )
            discovery = WebDiscovery(settings, storage)
            text = (
                "Đoạn thứ nhất có đủ từ riêng biệt để tạo truy vấn tìm kiếm công khai và đo thời gian phản hồi. "
                "Đoạn thứ hai bổ sung thuật ngữ khác nhau nhằm tạo thêm truy vấn tìm nguồn chạy song song. "
                "Đoạn thứ ba tiếp tục cung cấp nội dung độc lập để xác minh giới hạn chờ được áp dụng."
            )

            def slow_request(*_args, **_kwargs) -> dict:
                time.sleep(0.15)
                return {"results": []}

            started = time.monotonic()
            with patch.object(WebDiscovery, "_json_request", side_effect=slow_request):
                result = discovery.discover_and_index(text, organization_id=1)
            self.assertLess(time.monotonic() - started, 0.12)
            self.assertIn("Đã dừng chờ", result["message"])

    def test_global_deadline_skips_fallback_after_primary_uses_budget(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            storage = Storage(root / "test.db")
            settings = replace(
                settings_for(root, tavily_api_key="tavily", exa_api_key="exa"),
                web_discovery_time_budget_seconds=0.03,
            )
            discovery = WebDiscovery(settings, storage)
            text = "Nội dung đủ dài để xác minh deadline chung sẽ dừng fallback chậm và trả kết quả hiện có sớm."

            def slow_primary(*_args, **_kwargs):
                time.sleep(0.04)
                return DiscoveryResult("tavily", True, True, ["primary"], 0, 0, "Primary finished.", [])

            with (
                patch.object(discovery, "_tavily", side_effect=slow_primary),
                patch.object(discovery, "_exa") as exa,
            ):
                result = discovery.discover_and_index(text, organization_id=1)
            self.assertEqual(result["provider"], "tavily")
            exa.assert_not_called()


class CitationWritingAssistantTest(unittest.TestCase):
    def test_revision_uses_official_gemini_model_structured_output_and_citation_guardrails(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            assistant = CitationWritingAssistant(settings_for(root, gemini_api_key="gemini-test-key"))
            captured: dict = {}
            text = (
                "Liêm chính học thuật là nền tảng quan trọng của giáo dục đáng tin cậy và cần được "
                "thực hành bằng cách ghi nhận nguồn tham khảo rõ ràng trong từng bài viết của người học."
            )
            report = {
                "sources": [{"id": 1, "title": "Sổ tay học thuật", "url": "https://example.org/source"}],
                "matchedSegments": [
                    {
                        "text": "Liêm chính học thuật là nền tảng quan trọng của giáo dục đáng tin cậy.",
                        "source": {"id": 1},
                    }
                ],
                "webDiscovery": {
                    "regionalCoverage": {
                        "searchedRegions": 5,
                        "evidenceRegions": 2,
                        "totalRegions": 5,
                    }
                },
            }

            def fake_request(url, payload, *, headers, timeout):
                captured.update({"url": url, "payload": payload, "headers": headers, "timeout": timeout})
                return {
                    "candidates": [
                        {
                            "content": {
                                "parts": [
                                    {
                                        "text": json.dumps(
                                            {
                                                "revision": (
                                                    "Liêm chính học thuật góp phần xây dựng môi trường giáo dục đáng "
                                                    "tin cậy [Nguồn 1]. Người học cần ghi nhận rõ tài liệu tham khảo "
                                                    "trong quá trình hoàn thiện từng bài viết của mình."
                                                ),
                                                "editorNotes": ["Đã bổ sung vị trí dẫn nguồn."],
                                                "citationNotes": [
                                                    {
                                                        "marker": "[Nguồn 1]",
                                                        "sourceTitle": "Sổ tay học thuật",
                                                        "sourceUrl": "https://example.org/source",
                                                        "reason": "Ý tưởng được kế thừa từ nguồn đối chiếu.",
                                                    }
                                                ],
                                            }
                                        )
                                    }
                                ]
                            }
                        }
                    ]
                }

            with patch.object(assistant, "_json_request", side_effect=fake_request):
                result = assistant.revise(text, report)
            self.assertEqual(
                captured["url"],
                "https://generativelanguage.googleapis.com/v1beta/models/gemini-3-flash-preview:generateContent",
            )
            self.assertEqual(captured["headers"]["x-goog-api-key"], "gemini-test-key")
            self.assertEqual(captured["payload"]["generationConfig"]["thinkingConfig"]["thinkingLevel"], "high")
            self.assertIn("Do not help evade plagiarism detection", captured["payload"]["contents"][0]["parts"][0]["text"])
            self.assertIn("PUBLIC-WEB REGIONAL EVIDENCE", captured["payload"]["contents"][0]["parts"][0]["text"])
            self.assertIn("Missing public-web evidence does not prove originality", captured["payload"]["contents"][0]["parts"][0]["text"])
            self.assertIn("[Nguồn 1]", result["revision"])
            self.assertEqual(result["mode"], "citation-guided-revision")


class TextDisplayTest(unittest.TestCase):
    def test_repairs_reversible_utf8_mojibake_and_preserves_valid_vietnamese(self) -> None:
        self.assertEqual(normalize_display_text("LiÃªm chÃ­nh há»c thuáº­t"), "Liêm chính học thuật")
        self.assertEqual(normalize_display_text("Liêm chính học thuật"), "Liêm chính học thuật")

    def test_similarity_rewards_ordered_phrases_over_scrambled_keywords(self) -> None:
        source = "quản trị dữ liệu giáo dục minh bạch trách nhiệm giải trình bảo vệ quyền riêng tư người học"
        copied_with_extra = (
            "nghiên cứu cho thấy quản trị dữ liệu giáo dục minh bạch trách nhiệm giải trình "
            "bảo vệ quyền riêng tư người học trong trường học hiện đại"
        )
        scrambled = (
            "người học dữ liệu riêng tư trách nhiệm giáo dục bảo vệ giải trình minh bạch quản trị quyền"
        )
        self.assertGreater(similarity(source, copied_with_extra), 0.85)
        self.assertGreater(similarity(source, copied_with_extra), similarity(source, scrambled))


class ResponseWriterTest(unittest.TestCase):
    def test_broken_pipe_is_treated_as_client_disconnect(self) -> None:
        from backend.server import AppRequestHandler

        handler = AppRequestHandler.__new__(AppRequestHandler)
        handler.send_response = Mock()
        handler.send_header = Mock()
        handler.end_headers = Mock(side_effect=BrokenPipeError)
        handler.close_connection = False
        handler._write_response(b"ok", 200, [])
        self.assertTrue(handler.close_connection)


class AnalysisJobManagerTest(unittest.TestCase):
    def test_job_progress_is_monotonic_and_protected_by_token(self) -> None:
        manager = AnalysisJobManager(max_workers=1, ttl_seconds=60)
        release = threading.Event()

        def work(progress):
            progress(25, "extracting", "Đang đọc tài liệu.")
            release.wait(timeout=2)
            progress(78, "matching", "Đang đối chiếu.")
            return {"ok": True}

        created = manager.create(work)
        status = None
        for _attempt in range(50):
            status = manager.get(created["jobId"], created["jobToken"])
            if status and status["progress"] >= 25:
                break
            time.sleep(0.01)
        self.assertIsNone(manager.get(created["jobId"], "wrong-token"))
        self.assertIsNotNone(status)
        self.assertGreaterEqual(status["progress"], 25)
        release.set()
        for _attempt in range(50):
            status = manager.get(created["jobId"], created["jobToken"])
            if status and status["status"] == "completed":
                break
            time.sleep(0.01)
        self.assertEqual(status["status"], "completed")
        self.assertEqual(status["progress"], 100)
        self.assertEqual(status["result"], {"ok": True})


class CrawlerTest(unittest.TestCase):
    def test_respects_robots_txt_and_indexes_allowed_pages(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            site = ThreadingHTTPServer(("127.0.0.1", 0), LocalSiteHandler)
            thread = threading.Thread(target=site.serve_forever, daemon=True)
            thread.start()
            try:
                root = Path(directory)
                storage = Storage(root / "test.db")
                crawler = Crawler(storage, settings_for(root, allow_private=True))
                base = f"http://127.0.0.1:{site.server_address[1]}"
                crawler.enqueue(f"{base}/allowed")
                result = crawler.crawl_queued(max_pages=3, max_depth=1)
                urls = {source["url"] for source in storage.list_sources()}
                self.assertEqual(result["indexed"], 2)
                self.assertEqual(result["skipped"], 1)
                self.assertIn(f"{base}/allowed", urls)
                self.assertIn(f"{base}/next", urls)
                self.assertNotIn(f"{base}/blocked", urls)
            finally:
                site.shutdown()
                site.server_close()

    def test_rejects_private_hosts_by_default(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            storage = Storage(Path(directory) / "test.db")
            crawler = Crawler(storage, settings_for(Path(directory)))
            with self.assertRaises(ValueError):
                crawler.enqueue("http://127.0.0.1/private")

    def test_enqueues_pages_from_xml_sitemap(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            site = ThreadingHTTPServer(("127.0.0.1", 0), LocalSiteHandler)
            thread = threading.Thread(target=site.serve_forever, daemon=True)
            thread.start()
            try:
                root = Path(directory)
                storage = Storage(root / "test.db")
                crawler = Crawler(storage, settings_for(root, allow_private=True))
                base = f"http://127.0.0.1:{site.server_address[1]}"
                result = crawler.enqueue_sitemap(f"{base}/sitemap.xml")
                self.assertEqual(result["sitemaps"], 1)
                self.assertEqual(result["queued"], 2)
                self.assertEqual(storage.stats()["crawl_queue"]["queued"], 2)
            finally:
                site.shutdown()
                site.server_close()

    def test_retries_temporary_server_errors_until_page_recovers(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            LocalSiteHandler.flaky_hits = 0
            site = ThreadingHTTPServer(("127.0.0.1", 0), LocalSiteHandler)
            thread = threading.Thread(target=site.serve_forever, daemon=True)
            thread.start()
            try:
                root = Path(directory)
                storage = Storage(root / "test.db")
                crawler = Crawler(storage, settings_for(root, allow_private=True))
                base = f"http://127.0.0.1:{site.server_address[1]}"
                crawler.enqueue(f"{base}/flaky")
                result = crawler.crawl_queued(max_pages=3, max_depth=0)
                operations = storage.crawl_operations()
                self.assertEqual(result["retryScheduled"], 2)
                self.assertEqual(result["indexed"], 1)
                self.assertEqual(operations["queue"]["indexed"], 1)
                self.assertEqual(operations["recent"][0]["attempts"], 3)
            finally:
                site.shutdown()
                site.server_close()

    def test_failed_url_can_be_requeued_manually(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            site = ThreadingHTTPServer(("127.0.0.1", 0), LocalSiteHandler)
            thread = threading.Thread(target=site.serve_forever, daemon=True)
            thread.start()
            try:
                root = Path(directory)
                storage = Storage(root / "test.db")
                crawler = Crawler(storage, settings_for(root, allow_private=True, max_attempts=2))
                base = f"http://127.0.0.1:{site.server_address[1]}"
                crawler.enqueue(f"{base}/always-fail")
                result = crawler.crawl_queued(max_pages=2, max_depth=0)
                self.assertEqual(result["retryScheduled"], 1)
                self.assertEqual(result["failed"], 1)
                self.assertEqual(storage.stats()["crawl_queue"]["failed"], 1)
                self.assertEqual(storage.requeue_failed_urls(), 1)
                self.assertEqual(storage.stats()["crawl_queue"]["queued"], 1)
            finally:
                site.shutdown()
                site.server_close()


class ApiTest(unittest.TestCase):
    def test_health_stats_and_analysis_endpoints(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            server = create_server(settings_for(Path(directory), port=0))
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            try:
                base = f"http://127.0.0.1:{server.server_address[1]}"
                health = self._json(f"{base}/api/health")
                stats = self._json(f"{base}/api/stats")
                search_status = self._json(f"{base}/api/search/status")
                reindex = self._json(f"{base}/api/search/reindex", method="POST", payload={})
                report = self._json(
                    f"{base}/api/analyze",
                    method="POST",
                    payload={
                        "text": (
                            "Liêm chính học thuật là nền tảng của một môi trường giáo dục đáng tin cậy. "
                            "Người học cần phân biệt rõ việc tham khảo ý tưởng, trích dẫn trực tiếp và sao "
                            "chép nội dung mà không ghi nhận nguồn."
                        )
                    },
                )
                self.assertTrue(health["ok"])
                self.assertEqual(health["searchBackend"], "sqlite-fts5")
                self.assertEqual(health["webDiscoveryStrategy"], "adaptive-fingerprint-v4")
                self.assertEqual(health["webDiscoveryLimits"]["queries"], 10)
                self.assertEqual(health["webDiscoveryLimits"]["parallelWorkers"], 10)
                self.assertEqual(health["webDiscoveryLimits"]["mode"], "fast")
                self.assertEqual(health["webDiscoveryLimits"]["timeBudgetSeconds"], 22.0)
                self.assertEqual(health["webDiscoveryLimits"]["thoroughTimeBudgetSeconds"], 55.0)
                self.assertEqual(health["webDiscoveryLimits"]["fallbackMinSources"], 8)
                self.assertEqual(health["webDiscoveryLimits"]["exaMaxQueries"], 3)
                self.assertEqual(health["webDiscoveryLimits"]["exaMode"], "instant")
                self.assertEqual(health["webDiscoveryLimits"]["websearchapiMaxQueries"], 1)
                self.assertEqual(health["webDiscoveryLimits"]["linkupMaxQueries"], 1)
                self.assertEqual(health["webDiscoveryLimits"]["linkupDepth"], "fast")
                self.assertEqual(health["webDiscoveryLimits"]["serperMaxQueries"], 1)
                self.assertEqual(health["webDiscoveryLimits"]["enrichmentMaxSources"], 2)
                self.assertEqual(health["webDiscoveryLimits"]["geminiModel"], "gemini-3-flash-preview")
                self.assertEqual(health["webDiscoveryLimits"]["geminiExpansionMaxQueries"], 3)
                self.assertEqual(health["webDiscoveryLimits"]["geminiTimeoutSeconds"], 4.0)
                self.assertEqual(health["webDiscoveryLimits"]["geminiRevisionTimeoutSeconds"], 45.0)
                self.assertEqual(health["webDiscoveryLimits"]["geminiRevisionMaxInputChars"], 30000)
                self.assertFalse(health["writingAssistant"]["enabled"])
                self.assertEqual(health["writingAssistant"]["mode"], "citation-guided-revision")
                self.assertEqual(health["webDiscoveryLimits"]["openaiModel"], "gpt-5-nano")
                self.assertEqual(health["webDiscoveryLimits"]["openaiExpansionMaxQueries"], 3)
                self.assertEqual(health["webDiscoveryLimits"]["openaiTimeoutSeconds"], 4.0)
                self.assertEqual(search_status["backend"], "sqlite-fts5")
                self.assertEqual(reindex["chunks"], 20)
                self.assertEqual(stats["sources"], 4)
                self.assertGreater(report["percent"], 70)
                self.assertTrue(report["reportId"])
            finally:
                server.shutdown()
                server.server_close()

    def test_citation_revision_job_scans_original_and_proposal_before_returning(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            server = create_server(settings_for(Path(directory), port=0, gemini_api_key="gemini-test-key"))
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            try:
                base = f"http://127.0.0.1:{server.server_address[1]}"
                text = (
                    "Liêm chính học thuật là nền tảng của một môi trường giáo dục đáng tin cậy. "
                    "Người học cần phân biệt rõ việc tham khảo ý tưởng, trích dẫn trực tiếp và sao "
                    "chép nội dung mà không ghi nhận nguồn."
                )
                revision = (
                    "Trong môi trường giáo dục đáng tin cậy, người học cần ghi nhận rõ nguồn tham khảo "
                    "và phân biệt nội dung trích dẫn trực tiếp với ý tưởng được kế thừa [Nguồn 1]."
                )
                web_result = {
                    "provider": "tavily",
                    "enabled": True,
                    "externalProcessing": True,
                    "queries": ["liêm chính học thuật"],
                    "indexed": 1,
                    "skipped": 0,
                    "message": "Đã rà nguồn web công khai.",
                    "verificationMode": "thorough",
                    "sources": [],
                }
                proposal = {
                    "revision": revision,
                    "editorNotes": ["Đã thêm vị trí cần dẫn nguồn."],
                    "citationNotes": [],
                    "model": "gemini-3-flash-preview",
                    "mode": "citation-guided-revision",
                    "notice": "Đây là bản đề xuất.",
                }
                with (
                    patch.object(server.context.web_discovery, "discover_and_index", return_value=web_result) as discover,  # type: ignore[attr-defined]
                    patch.object(server.context.writing_assistant, "revise", return_value=proposal) as revise,  # type: ignore[attr-defined]
                ):
                    created = self._json(
                        f"{base}/api/analysis-jobs",
                        method="POST",
                        payload={"kind": "citation-revision", "text": text, "enableWebSearch": True},
                    )
                    result = self._poll_job(base, created)["result"]
                audit = self._json(f"{base}/api/audit")
                self.assertEqual(discover.call_count, 2)
                self.assertTrue(all(call.kwargs["thorough"] for call in discover.call_args_list))
                self.assertTrue(all(call.kwargs["max_results"] == 20 for call in discover.call_args_list))
                self.assertEqual(revise.call_count, 1)
                self.assertEqual(result["revision"], revision)
                self.assertTrue(result["externalWebVerification"])
                self.assertIn("verificationReport", result)
                self.assertIn(
                    "writing_assistant.citation_revision",
                    {event["action"] for event in audit["events"]},
                )
            finally:
                server.shutdown()
                server.server_close()

    def test_docx_upload_extracts_text_and_metadata(self) -> None:
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")

        with tempfile.TemporaryDirectory() as directory:
            server = create_server(settings_for(Path(directory), port=0))
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            try:
                document = Document()
                document.core_properties.author = "Minh Chung QA"
                document.add_paragraph(
                    "Liêm chính học thuật là nền tảng của một môi trường giáo dục đáng tin cậy. "
                    "Người học cần phân biệt rõ việc tham khảo ý tưởng, trích dẫn trực tiếp và sao "
                    "chép nội dung mà không ghi nhận nguồn."
                )
                stream = io.BytesIO()
                document.save(stream)
                report = self._json(
                    f"http://127.0.0.1:{server.server_address[1]}/api/analyze-upload",
                    method="POST",
                    payload={
                        "filename": "kiem-tra.docx",
                        "contentBase64": base64.b64encode(stream.getvalue()).decode("ascii"),
                    },
                )
                self.assertGreater(report["percent"], 70)
                self.assertEqual(report["documentMetadata"]["author"], "Minh Chung QA")
            finally:
                server.shutdown()
                server.server_close()

    def test_large_docx_upload_job_accepts_file_over_old_ten_megabyte_limit(self) -> None:
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")

        with tempfile.TemporaryDirectory() as directory:
            server = create_server(settings_for(Path(directory), port=0))
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            try:
                document = Document()
                document.core_properties.author = "Minh Chung Large DOCX QA"
                document.add_paragraph(
                    "Liêm chính học thuật là nền tảng của một môi trường giáo dục đáng tin cậy. "
                    "Người học cần phân biệt rõ việc tham khảo ý tưởng, trích dẫn trực tiếp và sao "
                    "chép nội dung mà không ghi nhận nguồn."
                )
                compact_stream = io.BytesIO()
                document.save(compact_stream)
                expanded_stream = io.BytesIO()
                with zipfile.ZipFile(io.BytesIO(compact_stream.getvalue())) as source_archive:
                    with zipfile.ZipFile(expanded_stream, "w") as destination_archive:
                        for item in source_archive.infolist():
                            destination_archive.writestr(item, source_archive.read(item.filename))
                        destination_archive.writestr(
                            "word/media/qa-padding.bin",
                            os.urandom(10_500_000),
                            compress_type=zipfile.ZIP_STORED,
                        )
                content = expanded_stream.getvalue()
                self.assertGreater(len(content), 10_000_000)
                base = f"http://127.0.0.1:{server.server_address[1]}"
                created = self._raw_json(
                    f"{base}/api/analysis-jobs/upload",
                    content=content,
                    headers={
                        "Content-Type": "application/octet-stream",
                        "X-Minh-Chung-Filename": "large-document.docx",
                    },
                )
                report = self._poll_job(base, created)["result"]
                self.assertGreater(report["percent"], 70)
                self.assertEqual(report["documentMetadata"]["author"], "Minh Chung Large DOCX QA")
            finally:
                server.shutdown()
                server.server_close()

    def test_opt_in_submission_becomes_internal_comparison_source_without_duplicates(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            server = create_server(settings_for(Path(directory), port=0))
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            try:
                base = f"http://127.0.0.1:{server.server_address[1]}"
                text = (
                    "Đây là bài nộp nội bộ thử nghiệm với nội dung riêng biệt để kiểm tra cơ chế đồng ý "
                    "lập chỉ mục cho những lần đối chiếu tiếp theo trong cùng một tổ chức giáo dục."
                )
                first = self._json(
                    f"{base}/api/analyze",
                    method="POST",
                    payload={"text": text, "indexForComparison": True},
                )
                second = self._json(
                    f"{base}/api/analyze",
                    method="POST",
                    payload={"text": text, "indexForComparison": True},
                )
                stats = self._json(f"{base}/api/stats")
                submissions = self._json(f"{base}/api/submissions")
                self.assertEqual(first["submissionId"], second["submissionId"])
                self.assertGreater(second["percent"], 90)
                self.assertEqual(second["sources"][0]["type"], "bài nộp nội bộ")
                self.assertEqual(stats["indexed_submissions"], 1)
                self.assertEqual(len(submissions["submissions"]), 1)
                deleted = self._json(f"{base}/api/submissions/{first['submissionId']}", method="DELETE")
                stats_after_delete = self._json(f"{base}/api/stats")
                self.assertTrue(deleted["ok"])
                self.assertEqual(stats_after_delete["indexed_submissions"], 0)
                self.assertEqual(stats_after_delete["sources"], 4)
            finally:
                server.shutdown()
                server.server_close()

    def test_operations_retry_and_source_version_endpoints(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            server = create_server(settings_for(Path(directory), port=0))
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            try:
                base = f"http://127.0.0.1:{server.server_address[1]}"
                source = self._json(
                    f"{base}/api/sources",
                    method="POST",
                    payload={
                        "title": "Nguồn API có phiên bản",
                        "url": "https://example.org/api-version",
                        "content": "Nguồn API có đủ nội dung để tạo phiên bản và kiểm tra endpoint lịch sử dữ liệu.",
                    },
                )
                versions = self._json(f"{base}/api/sources/{source['sourceId']}/versions")
                storage = server.context.storage  # type: ignore[attr-defined]
                storage.enqueue_url("https://example.org/failed")
                claimed = storage.claim_next_url()
                storage.finish_crawl(claimed["id"], "failed", "Lỗi thử nghiệm")
                operations = self._json(f"{base}/api/crawl/operations")
                retry = self._json(f"{base}/api/crawl/retry", method="POST", payload={"limit": 10})
                self.assertEqual(len(versions["versions"]), 1)
                self.assertEqual(operations["queue"]["failed"], 1)
                self.assertEqual(retry["requeued"], 1)
            finally:
                server.shutdown()
                server.server_close()

    def test_pdf_export_ocr_status_and_audit_endpoint(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            server = create_server(settings_for(Path(directory), port=0))
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            try:
                base = f"http://127.0.0.1:{server.server_address[1]}"
                report = self._json(
                    f"{base}/api/analyze",
                    method="POST",
                    payload={
                        "text": (
                            "Liêm chính học thuật là nền tảng của một môi trường giáo dục đáng tin cậy. "
                            "Người học cần phân biệt rõ việc tham khảo ý tưởng, trích dẫn trực tiếp và sao "
                            "chép nội dung mà không ghi nhận nguồn."
                        )
                    },
                )
                pdf = self._raw(f"{base}/api/reports/{report['reportId']}/pdf")
                ocr = self._json(f"{base}/api/ocr/status")
                audit = self._json(f"{base}/api/audit")
                self.assertTrue(pdf.startswith(b"%PDF"))
                self.assertIn("available", ocr)
                self.assertIn("reason", ocr)
                self.assertIn("report.export_pdf", {event["action"] for event in audit["events"]})
            finally:
                server.shutdown()
                server.server_close()

    def test_roles_and_organizations_isolate_private_sources(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            server = create_server(settings_for(Path(directory), port=0))
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            try:
                base = f"http://127.0.0.1:{server.server_address[1]}"
                private_text = (
                    "Cụm từ tenant biệt lập alpha beta gamma delta chỉ thuộc kho nội bộ của trường "
                    "Minh Chứng và được dùng để xác minh ranh giới dữ liệu giữa các tổ chức."
                )
                source = self._json(
                    f"{base}/api/sources",
                    method="POST",
                    payload={"title": "Nguồn riêng trường Minh Chứng", "content": private_text},
                    headers={"X-Minh-Chung-User": "demo-instructor"},
                )
                same_org = self._json(
                    f"{base}/api/analyze",
                    method="POST",
                    payload={"text": private_text},
                    headers={"X-Minh-Chung-User": "demo-student"},
                )
                with server.context.storage.connect() as connection:  # type: ignore[attr-defined]
                    cursor = connection.execute(
                        "INSERT INTO organizations(slug, name, created_at) VALUES (?, ?, ?)",
                        ("other-school", "Trường Khác", utc_now()),
                    )
                    other_organization_id = int(cursor.lastrowid)
                    connection.execute(
                        """
                        INSERT INTO users(organization_id, username, display_name, role, created_at)
                        VALUES (?, ?, ?, ?, ?)
                        """,
                        (other_organization_id, "other-student", "Sinh viên Trường Khác", "student", utc_now()),
                    )
                other_sources = self._json(
                    f"{base}/api/sources",
                    headers={"X-Minh-Chung-User": "other-student"},
                )
                other_org = self._json(
                    f"{base}/api/analyze",
                    method="POST",
                    payload={"text": private_text},
                    headers={"X-Minh-Chung-User": "other-student"},
                )
                with self.assertRaises(HTTPError) as error:
                    self._json(
                        f"{base}/api/crawl/operations",
                        headers={"X-Minh-Chung-User": "demo-student"},
                    )
                self.assertEqual(error.exception.code, 403)
                self.assertTrue(source["sourceId"])
                self.assertIn("Nguồn riêng trường Minh Chứng", {item["title"] for item in same_org["sources"]})
                self.assertNotIn(
                    "Nguồn riêng trường Minh Chứng",
                    {item["title"] for item in other_sources["sources"]},
                )
                self.assertNotIn("Nguồn riêng trường Minh Chứng", {item["title"] for item in other_org["sources"]})
            finally:
                server.shutdown()
                server.server_close()

    def test_web_discovery_runs_only_after_explicit_opt_in(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            server = create_server(settings_for(Path(directory), port=0))
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            try:
                base = f"http://127.0.0.1:{server.server_address[1]}"
                text = (
                    "Nội dung dùng để xác minh rằng quét nguồn web công khai chỉ chạy khi người dùng chủ động "
                    "bật lựa chọn gửi đoạn trích sang nhà cung cấp tìm kiếm bên ngoài để tìm nguồn liên quan."
                )
                web_result = {
                    "provider": "tavily",
                    "enabled": True,
                    "externalProcessing": True,
                    "queries": ["truy vấn thử nghiệm"],
                    "indexed": 1,
                    "skipped": 0,
                    "message": "Đã lập chỉ mục một nguồn.",
                    "sources": [],
                }
                with patch.object(server.context.web_discovery, "discover_and_index", return_value=web_result) as discover:  # type: ignore[attr-defined]
                    default_report = self._json(f"{base}/api/analyze", method="POST", payload={"text": text})
                    opted_in_report = self._json(
                        f"{base}/api/analyze",
                        method="POST",
                        payload={"text": text, "enableWebSearch": True},
                    )
                audit = self._json(f"{base}/api/audit")
                self.assertNotIn("webDiscovery", default_report)
                self.assertEqual(opted_in_report["webDiscovery"]["provider"], "tavily")
                discover.assert_called_once()
                self.assertIn("web_discovery.search", {event["action"] for event in audit["events"]})
            finally:
                server.shutdown()
                server.server_close()

    def test_public_mode_ignores_admin_header_and_does_not_store_guest_documents(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            server = create_server(settings_for(Path(directory), port=0, public_mode=True))
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            try:
                base = f"http://127.0.0.1:{server.server_address[1]}"
                admin_header = {"X-Minh-Chung-User": "demo-admin"}
                session = self._json(f"{base}/api/session", headers=admin_header)
                users = self._json(f"{base}/api/session/users", headers=admin_header)
                report = self._json(
                    f"{base}/api/analyze",
                    method="POST",
                    headers=admin_header,
                    payload={
                        "text": (
                            "Tài liệu khách công khai có đủ số lượng từ để tạo báo cáo thử nghiệm nhưng không "
                            "được lưu trên máy chủ hoặc đưa vào kho bài nộp nội bộ dù client gửi yêu cầu lưu."
                        ),
                        "saveReport": True,
                        "indexForComparison": True,
                    },
                )
                reports = self._json(f"{base}/api/reports", headers=admin_header)
                submissions = self._json(f"{base}/api/submissions", headers=admin_header)
                with self.assertRaises(HTTPError) as error:
                    self._json(f"{base}/api/crawl/operations", headers=admin_header)
                self.assertEqual(session["user"]["username"], "demo-student")
                self.assertEqual(session["user"]["role"], "student")
                self.assertEqual([user["username"] for user in users["users"]], ["demo-student"])
                self.assertFalse(report["dataRetention"]["stored"])
                self.assertNotIn("reportId", report)
                self.assertNotIn("submissionId", report)
                self.assertEqual(reports["reports"], [])
                self.assertEqual(submissions["submissions"], [])
                self.assertEqual(error.exception.code, 403)
            finally:
                server.shutdown()
                server.server_close()

    def test_password_mode_requires_cookie_and_persists_account_history(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            server = create_server(settings_for(Path(directory), port=0, auth_mode="password"))
            thread = threading.Thread(target=server.serve_forever, daemon=True)
            thread.start()
            try:
                base = f"http://127.0.0.1:{server.server_address[1]}"
                self.assertIn('id="auth-gate"', self._raw(f"{base}/").decode("utf-8"))
                with self.assertRaises(HTTPError) as error:
                    self._json(f"{base}/api/session", headers={"X-Minh-Chung-User": "demo-admin"})
                self.assertEqual(error.exception.code, 401)

                request = Request(
                    f"{base}/api/auth/register",
                    method="POST",
                    data=json.dumps(
                        {"username": "minh-test", "displayName": "Minh Test", "password": "mat-khau-123"}
                    ).encode("utf-8"),
                    headers={"Content-Type": "application/json"},
                )
                with urlopen(request, timeout=5) as response:
                    registered = json.loads(response.read().decode("utf-8"))
                    cookie = response.headers["Set-Cookie"].split(";", 1)[0]

                text = (
                    "Liêm chính học thuật là nền tảng của một môi trường giáo dục đáng tin cậy. "
                    "Người học cần phân biệt rõ việc tham khảo ý tưởng, trích dẫn trực tiếp và sao "
                    "chép nội dung mà không ghi nhận nguồn."
                )
                report = self._json(
                    f"{base}/api/analyze",
                    method="POST",
                    payload={"text": text},
                    headers={"Cookie": cookie},
                )
                history = self._json(f"{base}/api/reports", headers={"Cookie": cookie})
                self.assertEqual(registered["user"]["username"], "minh-test")
                self.assertTrue(report["reportId"])
                self.assertEqual(len(history["reports"]), 1)
            finally:
                server.shutdown()
                server.server_close()

    @staticmethod
    def _json(
        url: str,
        *,
        method: str = "GET",
        payload: dict | None = None,
        headers: dict[str, str] | None = None,
    ) -> dict:
        content = json.dumps(payload or {}).encode("utf-8") if payload is not None else None
        request = Request(
            url,
            method=method,
            data=content,
            headers={"Content-Type": "application/json", **(headers or {})},
        )
        with urlopen(request, timeout=5) as response:
            return json.loads(response.read().decode("utf-8"))

    @staticmethod
    def _raw(url: str, *, headers: dict[str, str] | None = None) -> bytes:
        request = Request(url, headers=headers or {})
        with urlopen(request, timeout=5) as response:
            return response.read()

    @staticmethod
    def _raw_json(
        url: str,
        *,
        content: bytes,
        headers: dict[str, str] | None = None,
    ) -> dict:
        request = Request(url, method="POST", data=content, headers=headers or {})
        with urlopen(request, timeout=15) as response:
            return json.loads(response.read().decode("utf-8"))

    @classmethod
    def _poll_job(cls, base: str, created: dict) -> dict:
        for _attempt in range(100):
            status = cls._json(
                f"{base}/api/analysis-jobs/{created['jobId']}",
                headers={"X-Minh-Chung-Job-Token": created["jobToken"]},
            )
            if status["status"] == "completed":
                return status
            if status["status"] == "failed":
                raise AssertionError(status.get("error") or status.get("message"))
            time.sleep(0.03)
        raise AssertionError("Analysis job did not complete in time.")


if __name__ == "__main__":
    unittest.main()
